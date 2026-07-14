"""
Generate mock data files (Excel + Word) for testing without Google Drive.
Run this script to create sample files in the mock_data/ directory.
"""

import sys
from datetime import datetime, timedelta
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_student_excel():
    """Create mock student tracking Excel file."""
    wb = openpyxl.Workbook()

    # ── Sheet 1: Lớp 10A1 ─────────────────────────────────
    ws1 = wb.active
    ws1.title = "Lớp 10A1"

    # Headers
    headers = ["STT", "Họ tên", "Lớp", "Môn", "Vấn đề", "Loại", "Deadline", "Trạng thái"]
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    # Data
    today = datetime.now()
    data = [
        [1, "Nguyễn Văn An", "10A1", "Toán", "Thiếu bài tập chương 2 - Hàm số", "nhắc nhở",
         (today + timedelta(days=2)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [2, "Trần Thị Bình", "10A1", "Toán", "Điểm TB cao nhất lớp - đề xuất khen thưởng", "khen thưởng",
         (today + timedelta(days=5)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [3, "Lê Văn Cường", "10A1", "Toán", "Vi phạm nội quy lớp 3 lần liên tiếp", "kỷ luật",
         (today + timedelta(days=1)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [4, "Phạm Thị Dung", "10A1", "Toán", "Thiếu bài kiểm tra giữa kỳ - cần kiểm tra bổ sung", "nhắc nhở",
         (today + timedelta(days=3)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [5, "Hoàng Văn Em", "10A1", "Toán", "Nghỉ học không phép 5 buổi", "kỷ luật",
         (today - timedelta(days=1)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [6, "Võ Thị Phương", "10A1", "Toán", "Đạt giải 3 cuộc thi Toán cấp trường", "khen thưởng",
         (today + timedelta(days=7)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [7, "Ngô Văn Giang", "10A1", "Toán", "Chưa nộp sổ liên lạc để ký", "nhắc nhở",
         (today + timedelta(days=0)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [8, "Đặng Thị Hoa", "10A1", "Toán", "Cần trao đổi với phụ huynh về tình hình học tập", "nhắc nhở",
         (today + timedelta(days=4)).strftime("%d/%m/%Y"), "Chưa xử lý"],
    ]

    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            if col_idx == 1:
                cell.alignment = Alignment(horizontal="center")

    # Adjust column widths
    ws1.column_dimensions["A"].width = 6
    ws1.column_dimensions["B"].width = 22
    ws1.column_dimensions["C"].width = 10
    ws1.column_dimensions["D"].width = 10
    ws1.column_dimensions["E"].width = 50
    ws1.column_dimensions["F"].width = 15
    ws1.column_dimensions["G"].width = 15
    ws1.column_dimensions["H"].width = 15

    # ── Sheet 2: Lớp 10A2 ─────────────────────────────────
    ws2 = wb.create_sheet("Lớp 10A2")

    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    data2 = [
        [1, "Bùi Văn Khải", "10A2", "Toán", "Thiếu vở bài tập", "nhắc nhở",
         (today + timedelta(days=2)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [2, "Lý Thị Lan", "10A2", "Toán", "Tiến bộ vượt bậc trong học tập", "khen thưởng",
         (today + timedelta(days=6)).strftime("%d/%m/%Y"), "Chưa xử lý"],
        [3, "Trương Văn Minh", "10A2", "Toán", "Thường xuyên đi học trễ", "kỷ luật",
         (today + timedelta(days=1)).strftime("%d/%m/%Y"), "Chưa xử lý"],
    ]

    for row_idx, row_data in enumerate(data2, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border

    ws2.column_dimensions["A"].width = 6
    ws2.column_dimensions["B"].width = 22
    ws2.column_dimensions["C"].width = 10
    ws2.column_dimensions["D"].width = 10
    ws2.column_dimensions["E"].width = 50
    ws2.column_dimensions["F"].width = 15
    ws2.column_dimensions["G"].width = 15
    ws2.column_dimensions["H"].width = 15

    # Save
    output_dir = Path(__file__).parent
    output_dir.mkdir(exist_ok=True)
    filepath = output_dir / "ThongKe_HocSinh.xlsx"
    wb.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


def create_task_word():
    """Create mock task planning Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("KẾ HOẠCH CÔNG VIỆC", level=1)

    doc.add_paragraph(
        f"Giáo viên: Nguyễn Nhật Bình\n"
        f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}\n"
        f"Năm học: 2026-2027"
    )

    doc.add_heading("I. Công việc cần thực hiện", level=2)

    # Create table
    today = datetime.now()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ["Công việc", "Mô tả", "Deadline", "Ưu tiên"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Data
    tasks = [
        [
            "Nộp kế hoạch giảng dạy HK1",
            "Gửi tổ trưởng bộ môn trước khi họp tổ",
            (today + timedelta(days=5)).strftime("%d/%m/%Y"),
            "Cao",
        ],
        [
            "Họp phụ huynh đầu năm lớp 10A1",
            "Chuẩn bị slide, phiếu thông tin HS",
            (today + timedelta(days=10)).strftime("%d/%m/%Y"),
            "Cao",
        ],
        [
            "Soạn đề kiểm tra 15 phút",
            "Chương 1 - Hàm số bậc hai, Toán 10",
            (today + timedelta(days=14)).strftime("%d/%m/%Y"),
            "Trung bình",
        ],
        [
            "Hoàn thành sổ chủ nhiệm tháng 8",
            "Cập nhật thông tin HS, điền nhận xét",
            (today + timedelta(days=20)).strftime("%d/%m/%Y"),
            "Trung bình",
        ],
        [
            "Đăng ký bồi dưỡng chuyên môn",
            "Khóa tập huấn SGK mới 2026",
            (today + timedelta(days=3)).strftime("%d/%m/%Y"),
            "Cao",
        ],
        [
            "Nộp hồ sơ đánh giá giáo viên",
            "Minh chứng BDTX, sáng kiến kinh nghiệm",
            (today + timedelta(days=25)).strftime("%d/%m/%Y"),
            "Thấp",
        ],
    ]

    for task_data in tasks:
        row_cells = table.add_row().cells
        for i, value in enumerate(task_data):
            row_cells[i].text = value

    doc.add_paragraph("")  # spacing
    doc.add_heading("II. Lịch họp", level=2)

    # Meeting table
    meeting_table = doc.add_table(rows=1, cols=4)
    meeting_table.style = "Table Grid"

    headers2 = ["Công việc", "Mô tả", "Deadline", "Ưu tiên"]
    header_cells2 = meeting_table.rows[0].cells
    for i, header in enumerate(headers2):
        header_cells2[i].text = header
        for paragraph in header_cells2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    meetings = [
        [
            "Họp hội đồng sư phạm",
            "Triển khai kế hoạch năm học mới",
            (today + timedelta(days=2)).strftime("%d/%m/%Y"),
            "Cao",
        ],
        [
            "Họp tổ bộ môn Toán",
            "Phân công chuyên môn, thống nhất PPCT",
            (today + timedelta(days=7)).strftime("%d/%m/%Y"),
            "Cao",
        ],
        [
            "Sinh hoạt chuyên đề cụm",
            "Đổi mới PPDH theo hướng phát triển năng lực",
            (today + timedelta(days=15)).strftime("%d/%m/%Y"),
            "Trung bình",
        ],
    ]

    for meeting in meetings:
        row_cells = meeting_table.add_row().cells
        for i, value in enumerate(meeting):
            row_cells[i].text = value

    # Save
    output_dir = Path(__file__).parent
    filepath = output_dir / "KeHoach_CongViec.docx"
    doc.save(filepath)
    print(f"✅ Created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("📁 Generating mock data files...\n")
    create_student_excel()
    create_task_word()
    print("\n🎉 Done! Mock files created in mock_data/")
