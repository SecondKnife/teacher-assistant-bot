"""
Document parser — extracts structured data from Excel (.xlsx) and Word (.docx) files.

Supports flexible column name mapping (Vietnamese headers).
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# ─── Column Name Mapping ─────────────────────────────────
# Maps Vietnamese column headers to standardized field names.
# Each field has a list of possible header variations.

STUDENT_COLUMN_MAP = {
    "stt": ["stt", "số thứ tự", "tt", "#"],
    "name": ["họ tên", "ho ten", "tên", "ten", "họ và tên", "ho va ten", "student"],
    "class_name": ["lớp", "lop", "class", "mã lớp"],
    "subject": ["môn", "mon", "môn học", "mon hoc", "subject"],
}

TASK_COLUMN_MAP = {
    "title": [
        "vấn đề", "van de", "công việc", "cong viec",
        "nội dung", "noi dung", "mô tả", "mo ta",
        "task", "title", "việc cần làm",
    ],
    "category": [
        "loại", "loai", "phân loại", "phan loai",
        "category", "type", "hình thức",
    ],
    "deadline": [
        "deadline", "hạn", "han", "ngày", "ngay",
        "thời hạn", "thoi han", "hạn chót", "due",
    ],
    "priority": [
        "ưu tiên", "uu tien", "priority", "mức độ", "muc do",
    ],
    "status": [
        "trạng thái", "trang thai", "status", "tình trạng", "tinh trang",
    ],
    "description": [
        "ghi chú", "ghi chu", "chi tiết", "chi tiet",
        "description", "note", "mô tả chi tiết",
    ],
}

# Category mapping from Vietnamese to enum values
CATEGORY_MAP = {
    "nhắc nhở": "nhắc nhở",
    "nhac nho": "nhắc nhở",
    "reminder": "nhắc nhở",
    "khen thưởng": "khen thưởng",
    "khen thuong": "khen thưởng",
    "reward": "khen thưởng",
    "kỷ luật": "kỷ luật",
    "ky luat": "kỷ luật",
    "discipline": "kỷ luật",
    "công việc": "công việc",
    "cong viec": "công việc",
    "task": "công việc",
}

PRIORITY_MAP = {
    "cao": "cao",
    "high": "cao",
    "trung bình": "trung bình",
    "trung binh": "trung bình",
    "medium": "trung bình",
    "tb": "trung bình",
    "thấp": "thấp",
    "thap": "thấp",
    "low": "thấp",
}


# ─── Helper Functions ─────────────────────────────────────


def _normalize_header(header: str) -> str:
    """Normalize a column header for matching."""
    return str(header).strip().lower()


def _find_column(df_columns: list[str], field_variants: list[str]) -> Optional[str]:
    """Find the actual column name in DataFrame that matches any of the variants."""
    normalized = {_normalize_header(col): col for col in df_columns}
    for variant in field_variants:
        if variant in normalized:
            return normalized[variant]
    return None


def _parse_date(value) -> Optional[datetime]:
    """Parse various date formats into datetime."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None

    if isinstance(value, datetime):
        return value

    if isinstance(value, pd.Timestamp):
        return value.to_pydatetime()

    # Try common Vietnamese date formats
    date_str = str(value).strip()
    formats = [
        "%d/%m/%Y",
        "%d-%m-%Y",
        "%Y-%m-%d",
        "%d/%m/%y",
        "%d-%m-%y",
        "%d.%m.%Y",
    ]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue

    logger.warning(f"Could not parse date: {value}")
    return None


def _parse_category(value) -> str:
    """Map category string to standardized value."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "công việc"
    key = str(value).strip().lower()
    return CATEGORY_MAP.get(key, "khác")


def _parse_priority(value) -> str:
    """Map priority string to standardized value."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "trung bình"
    key = str(value).strip().lower()
    return PRIORITY_MAP.get(key, "trung bình")


# ─── Data Structures ─────────────────────────────────────


class ParsedRecord:
    """A single parsed record from a document."""

    def __init__(
        self,
        title: str,
        student_name: str = None,
        class_name: str = None,
        subject: str = None,
        category: str = "công việc",
        deadline: datetime = None,
        priority: str = "trung bình",
        status: str = "pending",
        description: str = None,
        source_file: str = None,
    ):
        self.title = title
        self.student_name = student_name
        self.class_name = class_name
        self.subject = subject
        self.category = category
        self.deadline = deadline
        self.priority = priority
        self.status = status
        self.description = description
        self.source_file = source_file

    def __repr__(self):
        return f"<ParsedRecord(title='{self.title[:40]}', student='{self.student_name}')>"


class ParseResult:
    """Result of parsing a document."""

    def __init__(self, file_name: str):
        self.file_name = file_name
        self.records: list[ParsedRecord] = []
        self.errors: list[str] = []
        self.warnings: list[str] = []

    @property
    def success(self) -> bool:
        return len(self.errors) == 0

    @property
    def record_count(self) -> int:
        return len(self.records)

    def __repr__(self):
        return (
            f"<ParseResult(file='{self.file_name}', "
            f"records={self.record_count}, errors={len(self.errors)})>"
        )


# ─── Excel Parser ─────────────────────────────────────────


def parse_excel(file_path: str | Path) -> ParseResult:
    """
    Parse an Excel file containing student/task data.

    Supports multiple sheets — each sheet is parsed independently.
    """
    file_path = Path(file_path)
    result = ParseResult(file_name=file_path.name)

    if not file_path.exists():
        result.errors.append(f"File not found: {file_path}")
        return result

    try:
        xl = pd.ExcelFile(file_path)
        for sheet_name in xl.sheet_names:
            logger.info(f"Parsing sheet: {sheet_name}")
            df = xl.parse(sheet_name)

            if df.empty:
                result.warnings.append(f"Sheet '{sheet_name}' is empty")
                continue

            _parse_dataframe(df, result, file_path.name, sheet_name)

    except Exception as e:
        result.errors.append(f"Error reading Excel file: {e}")
        logger.exception(f"Failed to parse Excel: {file_path}")

    logger.info(f"Parsed {file_path.name}: {result.record_count} records")
    return result


# ─── Word Parser ──────────────────────────────────────────


def parse_word(file_path: str | Path) -> ParseResult:
    """
    Parse a Word (.docx) file containing tables with student/task data.

    Extracts all tables from the document.
    """
    file_path = Path(file_path)
    result = ParseResult(file_name=file_path.name)

    if not file_path.exists():
        result.errors.append(f"File not found: {file_path}")
        return result

    try:
        doc = DocxDocument(str(file_path))

        if not doc.tables:
            result.warnings.append("No tables found in Word document")
            return result

        for i, table in enumerate(doc.tables):
            logger.info(f"Parsing table {i + 1}/{len(doc.tables)}")
            df = _table_to_dataframe(table)

            if df.empty:
                result.warnings.append(f"Table {i + 1} is empty")
                continue

            _parse_dataframe(df, result, file_path.name, f"Table {i + 1}")

    except Exception as e:
        result.errors.append(f"Error reading Word file: {e}")
        logger.exception(f"Failed to parse Word: {file_path}")

    logger.info(f"Parsed {file_path.name}: {result.record_count} records")
    return result


def _table_to_dataframe(table) -> pd.DataFrame:
    """Convert a python-docx table to a pandas DataFrame."""
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    if not data:
        return pd.DataFrame()

    # First row as headers
    headers = data[0]
    rows = data[1:]

    return pd.DataFrame(rows, columns=headers)


# ─── Shared Parsing Logic ────────────────────────────────


def _parse_dataframe(
    df: pd.DataFrame,
    result: ParseResult,
    file_name: str,
    sheet_name: str = "",
):
    """Parse a DataFrame into ParsedRecords."""
    columns = list(df.columns)

    # Map columns
    name_col = _find_column(columns, STUDENT_COLUMN_MAP["name"])
    class_col = _find_column(columns, STUDENT_COLUMN_MAP["class_name"])
    subject_col = _find_column(columns, STUDENT_COLUMN_MAP["subject"])
    title_col = _find_column(columns, TASK_COLUMN_MAP["title"])
    category_col = _find_column(columns, TASK_COLUMN_MAP["category"])
    deadline_col = _find_column(columns, TASK_COLUMN_MAP["deadline"])
    priority_col = _find_column(columns, TASK_COLUMN_MAP["priority"])
    status_col = _find_column(columns, TASK_COLUMN_MAP["status"])
    desc_col = _find_column(columns, TASK_COLUMN_MAP["description"])

    if not title_col:
        result.warnings.append(
            f"Sheet/Table '{sheet_name}': Could not find task title column. "
            f"Available columns: {columns}"
        )
        return

    for idx, row in df.iterrows():
        try:
            title_val = row.get(title_col)
            if title_val is None or (isinstance(title_val, float) and pd.isna(title_val)):
                continue
            title = str(title_val).strip()
            if not title:
                continue

            record = ParsedRecord(
                title=title,
                student_name=(
                    str(row[name_col]).strip()
                    if name_col and pd.notna(row.get(name_col))
                    else None
                ),
                class_name=(
                    str(row[class_col]).strip()
                    if class_col and pd.notna(row.get(class_col))
                    else None
                ),
                subject=(
                    str(row[subject_col]).strip()
                    if subject_col and pd.notna(row.get(subject_col))
                    else None
                ),
                category=_parse_category(
                    row.get(category_col) if category_col else None
                ),
                deadline=_parse_date(
                    row.get(deadline_col) if deadline_col else None
                ),
                priority=_parse_priority(
                    row.get(priority_col) if priority_col else None
                ),
                description=(
                    str(row[desc_col]).strip()
                    if desc_col and pd.notna(row.get(desc_col))
                    else None
                ),
                source_file=file_name,
            )

            # Map status
            if status_col and pd.notna(row.get(status_col)):
                raw_status = str(row[status_col]).strip().lower()
                if raw_status in ("done", "hoàn thành", "xong", "đã xử lý"):
                    record.status = "done"
                elif raw_status in ("overdue", "trễ", "quá hạn"):
                    record.status = "overdue"

            result.records.append(record)

        except Exception as e:
            result.warnings.append(
                f"Row {idx + 2} in '{sheet_name}': {e}"
            )
            logger.warning(f"Error parsing row {idx + 2}: {e}")


# ─── Universal Parser ────────────────────────────────────


def parse_file(file_path: str | Path) -> ParseResult:
    """
    Parse any supported file type.
    Automatically detects format based on extension.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext in (".xlsx", ".xls"):
        return parse_excel(file_path)
    elif ext == ".docx":
        return parse_word(file_path)
    else:
        result = ParseResult(file_name=file_path.name)
        result.errors.append(f"Unsupported file format: {ext}")
        return result
